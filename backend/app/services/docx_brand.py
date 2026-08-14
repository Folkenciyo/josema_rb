"""Brand furniture for the Word exports: cover, running header, styles.

Word has no stylesheet to point at, so every brand decision has to be written
onto the document object itself. Keeping that here leaves the exporters to
describe the plan and nothing else.

Montserrat is named as the face; if the reader's machine does not have it Word
substitutes silently. That is a limitation of .docx as a format — the PDF export
is the one that always looks right.
"""

from pathlib import Path

from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table

BRAND_DIR = Path(__file__).resolve().parent.parent / "static" / "brand"

BRAND_RED = RGBColor(0x8B, 0x1E, 0x24)
INK = RGBColor(0x0B, 0x0B, 0x0D)
STEEL = RGBColor(0x6B, 0x6B, 0x6F)
BRAND_RED_HEX = "8B1E24"
ZEBRA_HEX = "FAFAFA"

DISPLAY_FONT = "Montserrat"


def _set_font(style, name: str) -> None:
    """python-docx only writes the ascii font; Word wants hAnsi named too."""
    style.font.name = name
    style.element.rPr.rFonts.set(qn("w:hAnsi"), name)


def apply_brand_styles(document: DocumentType) -> None:
    styles = document.styles

    normal = styles["Normal"]
    _set_font(normal, DISPLAY_FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x1C, 0x1C, 0x1E)

    for name, size, colour in (
        ("Heading 1", 16, INK),
        ("Heading 2", 12, BRAND_RED),
        ("Heading 3", 10, STEEL),
    ):
        style = styles[name]
        _set_font(style, DISPLAY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = colour

    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.6)


def _page_number_field(paragraph) -> None:
    """Word numbers pages through a field code, not literal text."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, end):
        run._r.append(element)
    run.font.size = Pt(7)
    run.font.color.rgb = STEEL


def _footer_line(paragraph) -> None:
    credit = paragraph.add_run("JOSEMA RB · Strength · Discipline · Results     ")
    credit.font.size = Pt(7)
    credit.font.color.rgb = STEEL
    _page_number_field(paragraph)


def add_running_furniture(document: DocumentType, subtitle: str) -> None:
    """Logo at the top of every page, credit and page number at the foot.

    Page one is the exception: it already carries the banner, and a second
    lockup above it would collide with the artwork — the same reason the PDF
    suppresses its running header on `@page :first`.
    """
    section = document.sections[0]
    section.different_first_page_header_footer = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.add_run().add_picture(str(BRAND_DIR / "logo-wordmark.png"), width=Cm(3.6))
    caption = header.add_run(f"    {subtitle}")
    caption.font.size = Pt(7.5)
    caption.font.color.rgb = STEEL

    _footer_line(section.footer.paragraphs[0])
    _footer_line(section.first_page_footer.paragraphs[0])


def add_cover(
    document: DocumentType,
    title: str,
    meta: list[tuple[str, str]],
    notes: str | None,
) -> None:
    """The banner, the plan title, and the facts about it, on their own page."""
    banner = document.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner.add_run().add_picture(str(BRAND_DIR / "document-header.png"), width=Cm(17.4))

    document.add_paragraph()
    heading = document.add_paragraph(title, style="Heading 1")
    heading.paragraph_format.space_after = Pt(10)

    for label, value in meta:
        line = document.add_paragraph()
        key = line.add_run(f"{label.upper()}   ")
        key.font.size = Pt(7.5)
        key.bold = True
        key.font.color.rgb = STEEL
        line.add_run(value).font.size = Pt(10)
        line.paragraph_format.space_after = Pt(3)

    if notes:
        document.add_paragraph()
        note = document.add_paragraph(notes)
        note.paragraph_format.left_indent = Cm(0.4)

    # No page break: on a two-page plan a dedicated cover is mostly blank paper.
    document.add_paragraph()


def shade(cell, hex_colour: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shading)


def _horizontal_rules_only(table: Table) -> None:
    """Drop the grid's vertical lines, as in the PDF: rows read as bands."""
    borders = OxmlElement("w:tblBorders")
    for edge, style, colour in (
        ("top", "single", "E5E5E7"),
        ("bottom", "single", "0B0B0D"),
        ("insideH", "single", "E5E5E7"),
        ("left", "none", "auto"),
        ("right", "none", "auto"),
        ("insideV", "none", "auto"),
    ):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), style)
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), colour)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def brand_table(
    document: DocumentType, headers: list[str], widths_cm: list[float]
) -> Table:
    """A table whose head is the brand red, matching the PDF export.

    Widths have to be written in two places. Every cell carries its own, and
    the table grid carries them again — python-docx leaves the grid at equal
    shares, and under a fixed layout that grid is what Word actually lays out
    from, which is why the headings would otherwise break mid-word.
    """
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    _horizontal_rules_only(table)

    for grid_column, width in zip(
        table._tbl.tblGrid.gridCol_lst, widths_cm, strict=True
    ):
        grid_column.set(qn("w:w"), str(Cm(width).twips))

    for cell, header, width in zip(
        table.rows[0].cells, headers, widths_cm, strict=True
    ):
        cell.width = Cm(width)
        shade(cell, BRAND_RED_HEX)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header.upper())
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    return table


def add_row(table: Table, values: list[str]) -> None:
    """Append a body row, zebra-striped and set at table body size."""
    cells = table.add_row().cells
    striped = len(table.rows) % 2 == 1

    for cell, value, head in zip(cells, values, table.rows[0].cells, strict=True):
        cell.width = head.width
        if striped:
            shade(cell, ZEBRA_HEX)
        run = cell.paragraphs[0].add_run(value)
        run.font.size = Pt(8.5)
