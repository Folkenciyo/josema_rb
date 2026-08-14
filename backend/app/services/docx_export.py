import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from app.schemas.export import DietPlanDocument, TrainingPlanDocument
from app.services.docx_brand import (
    add_cover,
    add_row,
    add_running_furniture,
    apply_brand_styles,
    brand_table,
)
from app.services.meal_template_service import NUTRIENT_FIELDS

STATIC_IMAGES_DIR = (
    Path(__file__).resolve().parent.parent / "static" / "exercise-images"
)

# Abbreviated as in the PDF: eight nutrient columns leave ~1.5 cm each, and the
# full words wrap onto a second line at that width.
NUTRIENT_HEADERS = [
    "Kcal",
    "Prot.",
    "Hidr.",
    "Azúc.",
    "Grasa",
    "Sat.",
    "Fibra",
    "Sal",
]

# Column widths in cm, summing to the 17.4 cm of body between the A4 margins.
TRAINING_WIDTHS_CM = [5.6, 1.8, 2.2, 2.4, 5.4]
DIET_WIDTHS_CM = [3.2, 2.2, *[1.5] * len(NUTRIENT_HEADERS)]


def _plan_meta(
    document: TrainingPlanDocument | DietPlanDocument,
) -> list[tuple[str, str]]:
    meta = [("Cliente", document.client_name)]
    if document.start_date:
        meta.append(("Inicio", str(document.start_date)))
    if document.end_date:
        meta.append(("Fin", str(document.end_date)))
    return meta


def _new_document(
    plan: TrainingPlanDocument | DietPlanDocument,
    meta: list[tuple[str, str]],
) -> Document:
    doc = Document()
    apply_brand_styles(doc)
    add_running_furniture(doc, f"{plan.client_name} · {plan.plan_title}")
    add_cover(doc, plan.plan_title, meta, plan.plan_notes)
    return doc


def _muted(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def render_training_plan_docx(document: TrainingPlanDocument) -> bytes:
    doc = _new_document(document, _plan_meta(document))

    for week in document.weeks:
        doc.add_heading(f"Semana {week.week_number}", level=1)
        if week.notes:
            doc.add_paragraph(week.notes)
        for day in week.days:
            doc.add_heading(day.day_of_week_es, level=2)
            if not day.exercises:
                _muted(doc, "Descanso")
                continue

            table = brand_table(
                doc,
                ["Ejercicio", "Series", "Reps", "Descanso", "Notas"],
                TRAINING_WIDTHS_CM,
            )
            for exercise in day.exercises:
                add_row(
                    table,
                    [
                        exercise.name_es,
                        str(exercise.sets),
                        exercise.reps,
                        f"{exercise.rest_seconds}s" if exercise.rest_seconds else "",
                        exercise.notes or "",
                    ],
                )

                if exercise.image_path:
                    image_path = STATIC_IMAGES_DIR / exercise.image_path
                    if image_path.exists():
                        doc.add_picture(str(image_path), width=Inches(1.5))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_diet_plan_docx(document: DietPlanDocument) -> bytes:
    meta = _plan_meta(document)
    if document.daily_calories_target:
        targets = [f"{document.daily_calories_target} kcal"]
        if document.daily_protein_g:
            targets.append(f"{document.daily_protein_g} g proteína")
        if document.daily_carbs_g:
            targets.append(f"{document.daily_carbs_g} g hidratos")
        if document.daily_fat_g:
            targets.append(f"{document.daily_fat_g} g grasa")
        meta.append(("Objetivo diario", " · ".join(targets)))

    doc = _new_document(document, meta)

    for week in document.weeks:
        doc.add_heading(f"Semana {week.week_number}", level=1)
        if week.notes:
            doc.add_paragraph(week.notes)
        for day in week.days:
            heading = day.day_of_week_es
            if day.menu_name:
                heading += f" — {day.menu_name}"
            doc.add_heading(heading, level=2)

            if not day.meals:
                _muted(doc, "Sin menú asignado")
                continue

            for meal in day.meals:
                meal_heading = meal.name
                if meal.time_of_day:
                    meal_heading += f" ({meal.time_of_day})"
                doc.add_heading(meal_heading, level=3)

                table = brand_table(
                    doc, ["Alimento", "Cantidad", *NUTRIENT_HEADERS], DIET_WIDTHS_CM
                )
                for item in meal.items:
                    values = [item.food_name, item.quantity_label or ""]
                    for field in NUTRIENT_FIELDS:
                        value = getattr(item, field)
                        values.append(str(value) if value is not None else "")
                    add_row(table, values)

                totals = ["Total de la comida", ""]
                totals += [
                    str(getattr(meal.totals, field)) for field in NUTRIENT_FIELDS
                ]
                add_row(table, totals)
                for cell in table.rows[-1].cells:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

            if day.totals:
                doc.add_paragraph(
                    f"Total del día: {day.totals.calories} kcal · "
                    f"{day.totals.protein_g} g proteína · "
                    f"{day.totals.carbs_g} g hidratos "
                    f"({day.totals.sugars_g} g azúcares) · "
                    f"{day.totals.fat_g} g grasa "
                    f"({day.totals.saturated_fat_g} g saturadas) · "
                    f"{day.totals.fiber_g} g fibra · {day.totals.salt_g} g sal"
                )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
